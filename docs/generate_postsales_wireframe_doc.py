"""Generate GA Post-Sales wireframe Word document."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUT = r"c:\Users\HP\OneDrive\Projects\API\Cursor\GA_Golden_Abodes_Platform\docs\GA_Post_Sales_App_Wireframe.docx"


def mono(doc, text, size=9):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(3)
    return p


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Cover
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("GA Post Sales Operations\n")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x18, 0x5F, 0xA5)
    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = s.add_run("Application Wireframe Document\n(for IT / Business review)")
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x5F, 0x5E, 0x5A)
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"\nGolden Abodes Platform  ·  Version 0.5.25  ·  {date.today().strftime('%d %B %Y')}")
    mr.font.size = Pt(11)
    doc.add_page_break()

    # TOC-style index
    heading(doc, "Document contents", 1)
    items = [
        "1. Application overview",
        "2. Global shell (header, navigation, sync banner)",
        "3. Screen wireframes",
        "   3.1  Dashboard",
        "   3.2  Allocation (admin)",
        "   3.3  My Tasks",
        "   3.4  Inventory setup",
        "   3.5  Units",
        "   3.6  Unit Pipeline (20-step SOP)",
        "   3.7  Documents",
        "   3.8  Demands / Collections",
        "   3.9  Loans",
        "   3.10 Tickets",
        "   3.11 Engineering milestones",
        "4. Navigation map",
        "5. 20-step pipeline reference",
        "6. Data & integration summary",
    ]
    for item in items:
        body(doc, item)
    doc.add_page_break()

    # 1 Overview
    heading(doc, "1. Application overview", 1)
    body(doc, (
        "GA Post Sales Operations is the sold-unit operating module within the Golden Abodes Platform. "
        "It is opened from the App Vault (Post Sales Operations) at route /app/post-sales. "
        "The module manages the full post-booking lifecycle: 20-step SOP pipeline, CLP collections, "
        "documents, home-loan tracking, customer tickets, and engineering milestone triggers."
    ))
    body(doc, "Primary users: CRM Executive (CX), Backend Executive, Accounts, Engineering, Management.")
    bullet(doc, "Hub entity: Sold Unit (project + phase + building + unit number)")
    bullet(doc, "Daily intake: CRM Excel upload on Units; collections maintained in Demands")
    bullet(doc, "Optional sync: Cashflow V1 sold units and collection rows on bootstrap")
    bullet(doc, "Role gate: post_sales app entitlement; Allocation tab requires admin password")
    doc.add_page_break()

    # 2 Global shell
    heading(doc, "2. Global shell", 1)
    body(doc, "Every screen shares the same top bar, navigation tabs, optional sync notice, and main content area.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  GA POST SALES OPERATIONS                          user@goldenabodes.com             │
│  Your working app for sold units, collections, pipeline & allocation               │
├──────────────────────────────────────────────────────────────────────────────────────┤
│ [Dashboard] [Allocation] [My Tasks] [Inventory] [Units] [Documents] [Demands]      │
│ [Loans] [Tickets] [Milestones]                                      [← Vault]       │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  ℹ Sync banner (on load): "142 units linked from Cashflow V1 · 380 collection rows  │
│    refreshed. Upload CLP & collections in Demands — Cashflow V1 reads from here."    │
├──────────────────────────────────────────────────────────────────────────────────────┤
│                                                                                      │
│                         << ACTIVE SCREEN CONTENT >>                                  │
│                                                                                      │
└──────────────────────────────────────────────────────────────────────────────────────┘
""")
    body(doc, "Shared filter bar (most list screens): Project ▼  Phase ▼  Building ▼  [Clear filters]")
    doc.add_page_break()

    # 3 Screens
    heading(doc, "3. Screen wireframes", 1)

    # 3.1 Dashboard
    heading(doc, "3.1 Dashboard", 2)
    body(doc, "Route: /app/post-sales  ·  Purpose: Operations KPIs and drill-down to breach units / tickets.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  Operations Dashboard                                                                │
│  [Project ▼] [Phase ▼] [Building ▼] [Clear]                                          │
├──────────────────────────────────────────────────────────────────────────────────────┤
│ ┌─────────────┐ ┌─────────────┐ ┌─────────────┐ ┌─────────────┐                      │
│ │ Active units│ │ SLA breaches│ │ Open tickets│ │ Outstanding │                      │
│ │     847     │ │      23     │ │      41     │ │  ₹12.4 Cr   │                      │
│ └─────────────┘ └─────────────┘ └─────────────┘ └─────────────┘                      │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  Collection progress                                                                 │
│  ₹8.2 Cr collected  ·  66% of ₹12.4 Cr demanded                                      │
│  [████████████████████░░░░░░░░░░]                                                    │
├───────────────────────────────┬──────────────────────────────────────────────────────┤
│  Units by project             │  Units by pipeline phase                             │
│  Paradise        ████ 312     │  ● Booking handoff        45                         │
│  Wakad GA        ███  198     │  ● Documentation & loan   120                        │
│  Golden HQ       ██   87      │  ● CLP demands            210                        │
├───────────────────────────────┴──────────────────────────────────────────────────────┤
│  SLA breach units (click → Unit Pipeline)                                            │
│  Paradise · 1002 — Sharma  ·  Step 12 overdue                                       │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  Open tickets list  ·  Pending engineering milestones                                │
└──────────────────────────────────────────────────────────────────────────────────────┘
""")

    # 3.2 Allocation
    heading(doc, "3.2 Allocation (admin)", 2)
    body(doc, "Route: /app/post-sales/allocation  ·  Password-gated admin workspace.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  Work allocation                                                                     │
│  [Enter admin password]  [Unlock]                                                    │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  Sub-tabs: [Work assignment] [Activity catalog]                                      │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  WORK ASSIGNMENT TAB                                                                 │
│  Filters: Project ▼ Phase ▼ Building ▼  Work type: All | CX | Backend              │
│  Bulk executives: CX [Priya ▼]  Backend [Raj ▼]  [Apply to selected]               │
│  Bulk steps: Kind [CX ▼]  Assignee [▼]  [Assign open steps]  [Auto-assign defaults]│
├──────────────────────────────────────────────────────────────────────────────────────┤
│  ☐ │ Unit  │ Customer │ Step │ Activity           │ CX exec │ Backend │ Pending  │
│  ☑ │ 1002  │ Sharma   │  12  │ CLP demand letter  │ Priya   │ Raj     │ ₹5.4L    │
│  ☐ │ 804   │ Patel    │   4  │ Home loan coord.   │ Priya   │ Amit    │ —        │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  ACTIVITY CATALOG TAB — edit 20-step SOP definitions, SLAs, checklists (no deploy)  │
│  Step # │ Name │ Phase │ Task kind │ SLA │ Checklist lines…  [Edit] [Save]           │
└──────────────────────────────────────────────────────────────────────────────────────┘
""")

    # 3.3 My Tasks
    heading(doc, "3.3 My Tasks", 2)
    body(doc, "Route: /app/post-sales/my-tasks  ·  Personal queue of open pipeline steps assigned to logged-in user.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  My Tasks                                                                            │
│  Open pipeline steps assigned to you — sorted by next action date, then due date.    │
│  Badges: [3 overdue] [5 due soon] [18 open]                                            │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  Tabs: [All (18)] [CX (12)] [Backend (6)]                                            │
│  [Project ▼] [Phase ▼] [Building ▼] [Clear]                                          │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  Unit   │ Step │ Phase        │ Activity              │ Due     │ SLA    │ Action  │
│  1002   │  12  │ CLP demands  │ CLP demand letter     │ 17 Jun  │ Overdue│ [Open]  │
│  804    │   4  │ Doc & loan   │ Home loan coordination│ 20 Jun  │ 2d left│ [Open]  │
│  1205   │   9  │ Agreement    │ Agreement execution   │ 22 Jun  │ OK     │ [Open]  │
└──────────────────────────────────────────────────────────────────────────────────────┘
""")

    # 3.4 Inventory
    heading(doc, "3.4 Inventory setup", 2)
    body(doc, "Route: /app/post-sales/inventory  ·  Master catalog Project → Phase → Building.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  Inventory setup                                                                     │
│  [+ Add project]  [Import from Cashflow V1]  [Push to Cashflow V1]                    │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  ▶ Paradise (PAD · Goa)                                                              │
│      ▶ Phase 1                                                                       │
│          · Tower A    [Edit] [Delete]                                                │
│          · Tower B    [Edit] [Delete]  [+ Add building]                              │
│      [+ Add phase]                                                                   │
│  ▶ Wakad GA (NBD · Wakad, Pune)                                                      │
└──────────────────────────────────────────────────────────────────────────────────────┘
""")

    # 3.5 Units
    heading(doc, "3.5 Units", 2)
    body(doc, "Route: /app/post-sales/units  ·  Sold-unit registry; CRM upload; Cashflow V1 sync.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  Sold units                                                                          │
│  [+ New unit]  [Upload CRM data]  [Sync from Cashflow V1]  [Preview sync]  [Purge]  │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  [Project ▼] [Phase ▼] [Building ▼]  Status ▼  Import batch ▼  [Clear]             │
│  SLA breaches: 12 units                                                              │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  Unit │ Customer │ Project  │ Step │ Status      │ Agreement │ CX      │ Backend   │
│  1002 │ Sharma   │ Paradise │  12  │ In progress │ ₹54L      │ Priya   │ Raj       │
│  804  │ Patel    │ Wakad GA │   4  │ In progress │ ₹42L      │ Priya   │ Amit      │
│       │          │          │      │             │           │         │ [Pipeline]│
└──────────────────────────────────────────────────────────────────────────────────────┘
""")

    # 3.6 Unit Pipeline
    heading(doc, "3.6 Unit Pipeline (20-step SOP)", 2)
    body(doc, "Route: /app/post-sales/units/:id  ·  Primary per-unit workspace.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  ← Units    UNIT 1002 · Paradise · Phase 2 · Tower B                                 │
│  Customer: Sharma · Booking: 12-Jan-24 · Total cost: ₹54,00,000                      │
│  CX: Priya · Backend: Raj · Entity: PAD · Funding: Home loan                         │
│  Overall: In progress · Current step: 12 · SLA breach: Yes                           │
├───────────────────────────────┬──────────────────────────────────────────────────────┤
│  PIPELINE STEPS (left rail)   │  STEP 12 — CLP demand letter issuance                │
│  ─ Booking handoff ─          │  Status: [Pending|In progress|Completed|Overdue]       │
│  ✓ 1 Booking confirm to CRM   │  Assignee: [Priya ▼]  Due: 17-Jun  Next action: —   │
│  ✓ 2 Update booking in CRM    │  Tabs: [Checklist][Documents][SOP details][Escalation]│
│  …                            │  ─ Checklist ─                                       │
│  ─ CLP demands ─              │  ☐ Demand letter drafted                             │
│  ● 12 CLP demand letter       │  ☐ Sent to customer                                  │
│  ○ 13 Payment follow-up       │  ☐ Copy filed in document vault                      │
│  ─ Possession ─               │  [Add comment]  [Set next action date]               │
│  ○ 14 Possession clearance    │  ─ Documents for this step ─                         │
│                               │  [Upload] CLP demand letter PDF                      │
└───────────────────────────────┴──────────────────────────────────────────────────────┘
""")

    # 3.7 Documents
    heading(doc, "3.7 Documents", 2)
    body(doc, "Route: /app/post-sales/documents  ·  Unit document vault with search and inline upload.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  Document vault                                                                      │
├───────────────────────────────┬──────────────────────────────────────────────────────┤
│  Select unit                  │  Documents for 1002 · Paradise · Sharma              │
│  Search units…                │  Search documents…  [+ Upload document]               │
│  · 1002 Paradise              │  Status: Received 12 · Verified 8 · Pending 5      │
│  · 804 Wakad GA               │  ─ Booking & KYC ─                                 │
│  · 1205 Golden HQ             │  Booking form      [Verified]  [Open]  [Upload]      │
│                               │  PAN — applicant   [Received]  [Open]  [Upload]      │
│                               │  ─ Agreement ─                                     │
│                               │  Draft agreement   [Pending]   —       [Upload]      │
│                               │  Registered copy   [—]           —       [Upload]      │
└───────────────────────────────┴──────────────────────────────────────────────────────┘
""")

    # 3.8 Demands
    heading(doc, "3.8 Demands / Collections", 2)
    body(doc, "Route: /app/post-sales/demands  ·  CLP collections ledger; canonical source for Cashflow V1 export.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  Demands & collections                                                               │
│  Cumulative as of today — agreement due/received only for CLP stages due ≤ today.    │
│  GST due/received from CRM GST column.                                               │
│  [Upload Excel] [Sync from Cashflow V1]  View: [By unit ●] [All milestones]          │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  KPI: Milestone rows │ Total due │ Received (%) │ Pending                            │
│  [All][Pending][Partial][Paid][Overdue]   Search unit, customer, milestone…           │
│  [Project ▼] [Phase ▼] [Building ▼]                                                  │
├──────────────────────────────────────────────────────────────────────────────────────┤
│ Unit │Cust. │Agmt due│Agmt recd│Agmt pend│GST due│GST recd│GST pend│Status│ ▶       │
│ 1002 │Sharma│ 32.4L  │  27.0L  │  5.4L   │ 6.7L  │ 1.35L  │ 5.35L  │Part. │ [▼]   │
│  └─ expanded milestone rows (sticky header):                                         │
│     Milestone      │Target│Actual│Agmt due│Agmt recd│Agmt pend│GST due│ [Pay]       │
│     Token          │ —    │ —    │  5.0L  │  5.0L   │    0    │   —   │             │
│     Slab complete  │01-Jun│01-Jun│ 10.8L  │  8.1L   │  2.7L   │   —   │ [Record]    │
│     GST (unit row) │ —    │ —    │   —    │    —    │    —    │ 6.7L  │             │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  Page totals row: Agreement pending ₹… · GST pending ₹…                              │
└──────────────────────────────────────────────────────────────────────────────────────┘
""")

    # 3.9 Loans
    heading(doc, "3.9 Loans", 2)
    body(doc, "Route: /app/post-sales/loans  ·  Home-loan tracker or self-funded contribution schedule.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  Loan & funding tracker                                                              │
├───────────────────────────────┬──────────────────────────────────────────────────────┤
│  Unit list                    │  Home loan — 1002 · Sharma                           │
│  1002  [Home loan] Paradise   │  Applied → Processing → Valuation → Sanctioned       │
│  804   [Self-funded] Wakad GA │  Bank: HDFC · RM: Kumar · Sanction: ₹43L            │
│  1205  [Home loan] Golden HQ  │  Disbursements: [table]  [Edit]                      │
│                               │  ─ OR Self-funded ─                                  │
│                               │  Own contribution schedule · Mark paid per row       │
└───────────────────────────────┴──────────────────────────────────────────────────────┘
""")

    # 3.10 Tickets
    heading(doc, "3.10 Tickets", 2)
    body(doc, "Route: /app/post-sales/tickets  ·  Customer queries, grievances, defects with SLA tracking.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  Customer tickets     Ack breaches: 3   Resolution breaches: 7                       │
├───────────────────────────────┬──────────────────────────────────────────────────────┤
│  Filter                       │  Ticket detail — TK-1042                             │
│  [All][Open][SLA breach]      │  Type: Grievance · Paradise · 1002                   │
│  [Query][Grievance][Defect]   │  Description: …                                      │
│  ─ list ─                     │  Activity log · Acknowledge · Resolve · Escalate     │
│  TK-1042 Grievance            │  SLA: Ack 24h · Resolution 7d                        │
│  Paradise · 1002              │                                                      │
│  [+ New ticket]               │                                                      │
└───────────────────────────────┴──────────────────────────────────────────────────────┘
""")

    # 3.11 Milestones
    heading(doc, "3.11 Engineering milestones", 2)
    body(doc, "Route: /app/post-sales/milestones  ·  Log construction CLP events; auto-create demands + Step 12 tasks.")
    mono(doc, """
┌──────────────────────────────────────────────────────────────────────────────────────┐
│  Engineering milestones                    [+ Log milestone]                         │
│  Pending: 2 · Triggered: 5 · Completed: 18                                           │
├──────────────────────────────────────────────────────────────────────────────────────┤
│  Project │ Tower │ Milestone      │ CLP % │ Logged   │ Units │ Status    │ Action  │
│  Paradise│ B     │ Slab complete  │  10%  │ 01-Jun   │  142  │ Triggered │ [Retry] │
│  Wakad GA│ A     │ Plinth complete│   5%  │ 15-May   │  138  │ Completed │         │
└──────────────────────────────────────────────────────────────────────────────────────┘
""")
    doc.add_page_break()

    # 4 Navigation map
    heading(doc, "4. Navigation map", 1)
    mono(doc, """
App Vault → Post Sales Operations (/app/post-sales)
│
├── Dashboard ..................... KPIs, breaches, collection progress
├── Allocation .................... Admin: bulk assign + activity catalog
├── My Tasks ...................... Personal pipeline queue (CX / Backend)
├── Inventory ..................... Project → Phase → Building catalog
├── Units ......................... Sold-unit list → Unit Pipeline
│   └── /units/:id ................ 20-step SOP (checklist, docs, escalation)
├── Documents ..................... Cross-unit document vault
├── Demands ....................... CLP collections (Agreement + GST split)
├── Loans ......................... Home loan / self-funded tracker
├── Tickets ....................... Customer service cases
└── Milestones .................... Engineering → demand + CLP letter trigger
""")

    # 5 Pipeline
    heading(doc, "5. Twenty-step pipeline reference", 1)
    phases = [
        ("Booking handoff", "Steps 1–3"),
        ("Documentation & loan", "Steps 4–5"),
        ("Agreement", "Steps 6–10"),
        ("Post-registration", "Step 11"),
        ("CLP demands", "Steps 12–13"),
        ("Servicing", "Step 14 (partial)"),
        ("Possession", "Step 14 clearance"),
        ("CHS formation", "Steps 16–18"),
        ("Post-handover", "Steps 19–20"),
    ]
    for phase, steps in phases:
        bullet(doc, f"{phase}: {steps}")

    body(doc, "Full step names, SLAs, checklists, and escalation rules are defined in the Activity catalog (Allocation tab) and server configuration.")

    # 6 Data summary
    heading(doc, "6. Data & integration summary", 1)
    body(doc, "Key entities: Customer, Unit, PipelineStep (×20 per unit), Demand (CLP + GST row), Document, LoanTracker, Ticket, ConstructionMilestone.")
    body(doc, "Integrations:")
    bullet(doc, "CRM Excel upload (Units): unit master + collection report blocks")
    bullet(doc, "Cashflow V1 bootstrap: sold units and collection rows on app load")
    bullet(doc, "Demands export: Cashflow V1 reads collections from Post-Sales")
    bullet(doc, "Engineering milestone log: creates demand rows + Step 12 tasks for matching tower units")

    body(doc, "")
    body(doc, "— End of wireframe document —")

    doc.save(OUT)
    print(f"Written: {OUT}")


if __name__ == "__main__":
    build()
